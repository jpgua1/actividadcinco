"""
api.py - Project Generator Agent Server

Servidor FastAPI que transforma documentación funcional/técnica en artefactos
completos para arranque de proyectos digitales.

Puerto: 8011
"""

import os
import json
import uuid
import threading
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List
import logging

from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse, Response
from pydantic import BaseModel, Field
from dotenv import load_dotenv

# Imports para exportación de documentos
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

# Cargar .env
load_dotenv()

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# FastAPI App
app = FastAPI(
    title="Project Generator Agent",
    description="Genera artefactos completos para arranque de proyectos digitales",
    version="1.0.0"
)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directorios
OUTPUT_DIR = Path("./outputs")
OUTPUT_DIR.mkdir(exist_ok=True)

# Almacén de jobs en memoria
jobs_store: Dict[str, Dict[str, Any]] = {}
jobs_lock = threading.Lock()


# ============================================================
# MODELOS DE DATOS
# ============================================================

class AnalyzeRequest(BaseModel):
    documentation: str = Field(..., description="Documentación funcional/técnica del proyecto")
    project_name: Optional[str] = Field(default="Proyecto Digital", description="Nombre del proyecto")
    context: Optional[str] = Field(default="", description="Contexto adicional del proyecto")


class JobResponse(BaseModel):
    job_id: str
    status: str
    message: str
    progress: Optional[float] = None
    created_at: str
    updated_at: str
    result: Optional[Dict[str, Any]] = None


# ============================================================
# HELPERS
# ============================================================

def update_job(job_id: str, status: str = None, message: str = None, progress: float = None, result: Any = None):
    """Actualiza el estado de un job de manera thread-safe"""
    with jobs_lock:
        if job_id in jobs_store:
            if status:
                jobs_store[job_id]["status"] = status
            if message:
                jobs_store[job_id]["message"] = message
            if progress is not None:
                jobs_store[job_id]["progress"] = progress
            if result is not None:
                jobs_store[job_id]["result"] = result
            jobs_store[job_id]["updated_at"] = datetime.now().isoformat()


def add_markdown_paragraph(doc, text: str, style: str = None) -> None:
    """
    Agrega un párrafo al documento Word convirtiendo markdown a formato.
    Soporta **negrita** y *cursiva*.
    """
    if not text:
        return
    
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    
    # Procesar texto con formato markdown
    remaining = text
    while remaining:
        # Buscar **negrita**
        if '**' in remaining:
            before, _, after = remaining.partition('**')
            if before:
                p.add_run(before)
            if '**' in after:
                bold_text, _, remaining = after.partition('**')
                run = p.add_run(bold_text)
                run.bold = True
            else:
                p.add_run(after)
                break
        # Buscar *cursiva*
        elif '*' in remaining:
            before, _, after = remaining.partition('*')
            if before:
                p.add_run(before)
            if '*' in after:
                italic_text, _, remaining = after.partition('*')
                run = p.add_run(italic_text)
                run.italic = True
            else:
                p.add_run(after)
                break
        else:
            p.add_run(remaining)
            break
    
    return p


# Importar orquestador
try:
    from AS_langgraphEnabler import ProjectGeneratorOrchestrator, ProjectConfig
    HAS_ORCHESTRATOR = True
    orchestrator = ProjectGeneratorOrchestrator()
    logger.info("✅ Orquestador LangGraph inicializado")
except Exception as e:
    logger.warning(f"⚠️ No se pudo cargar el orquestador: {e}")
    HAS_ORCHESTRATOR = False
    orchestrator = None


def generate_fallback_result(documentation: str, project_name: str) -> Dict[str, Any]:
    """Genera resultados básicos sin LLM"""
    return {
        "status": "completed",
        "architecture_diagram": """flowchart TD
    A[Cliente Web/Mobile] --> B[API Gateway]
    B --> C[Backend Services]
    C --> D[(Base de Datos)]
    C --> E[Cache Redis]
    C --> F[Message Queue]
    F --> G[Workers]""",
        "gantt_diagram": """gantt
    title Cronograma del Proyecto
    dateFormat YYYY-MM-DD
    section Análisis
    Levantamiento :a1, 2025-01-20, 10d
    Diseño técnico :a2, after a1, 5d
    section Desarrollo
    Sprint 1 - Core :a3, after a2, 15d
    Sprint 2 - Features :a4, after a3, 15d
    Sprint 3 - Integración :a5, after a4, 15d
    section QA
    Testing :a6, after a5, 10d
    UAT :a7, after a6, 5d""",
        "team_proposal": {
            "team_size": 5,
            "duration_months": 4,
            "roles": [
                {"role": "Tech Lead", "count": 1, "seniority": "Senior", "skills": ["Arquitectura", "Liderazgo técnico"], "dedication": "100%"},
                {"role": "Backend Developer", "count": 2, "seniority": "Mid-Senior", "skills": ["Python/Node.js", "APIs REST"], "dedication": "100%"},
                {"role": "Frontend Developer", "count": 1, "seniority": "Mid", "skills": ["React/Vue", "CSS"], "dedication": "100%"},
                {"role": "QA Engineer", "count": 1, "seniority": "Mid", "skills": ["Testing", "Automation"], "dedication": "100%"}
            ],
            "recommended_methodology": "Scrum",
            "sprint_duration": "2 semanas"
        },
        "epics": [
            {
                "id": "EPIC-001",
                "title": "Infraestructura y Setup",
                "description": "Configuración inicial del proyecto y CI/CD",
                "priority": "Alta",
                "estimated_sprints": 1,
                "stories": []
            },
            {
                "id": "EPIC-002",
                "title": "Módulo Core",
                "description": "Funcionalidades principales del sistema",
                "priority": "Alta",
                "estimated_sprints": 2,
                "stories": []
            },
            {
                "id": "EPIC-003",
                "title": "Integraciones",
                "description": "Integración con sistemas externos",
                "priority": "Media",
                "estimated_sprints": 1,
                "stories": []
            }
        ],
        "economic_proposal": {
            "summary": {
                "total_cost": 48000,
                "duration_months": 4,
                "monthly_cost": 12000
            },
            "breakdown": [
                {"category": "Desarrollo", "description": "Equipo de desarrollo", "cost": 40000, "recurrence": "Único"},
                {"category": "Infraestructura", "description": "Cloud y servicios", "cost": 4000, "recurrence": "Mensual"},
                {"category": "Licencias", "description": "Herramientas y software", "cost": 2000, "recurrence": "Único"},
                {"category": "QA", "description": "Testing y calidad", "cost": 2000, "recurrence": "Único"}
            ],
            "assumptions": [
                "Equipo dedicado al proyecto",
                "Infraestructura cloud (AWS/Azure/GCP)",
                "Metodología ágil con sprints de 2 semanas"
            ],
            "payment_milestones": [
                {"milestone": "Kickoff", "percentage": 20},
                {"milestone": "MVP", "percentage": 40},
                {"milestone": "Entrega Final", "percentage": 40}
            ]
        },
        "risks_prerequisites": {
            "risks": [
                {"id": "R001", "category": "Técnico", "description": "Complejidad técnica subestimada", "probability": "Media", "impact": "Alto", "mitigation": "Spike técnico en Sprint 0"},
                {"id": "R002", "category": "Recurso", "description": "Disponibilidad del cliente para validaciones", "probability": "Media", "impact": "Medio", "mitigation": "Definir agenda de demos fija"},
                {"id": "R003", "category": "Entrega", "description": "Cambios de alcance no controlados", "probability": "Alta", "impact": "Alto", "mitigation": "Proceso formal de gestión de cambios"}
            ],
            "prerequisites": [
                {"id": "P001", "category": "Acceso", "description": "Acceso a sistemas y repositorios del cliente", "owner": "Cliente", "deadline": "Antes del kickoff"},
                {"id": "P002", "category": "Documentación", "description": "Requerimientos funcionales detallados", "owner": "Cliente", "deadline": "Antes del kickoff"},
                {"id": "P003", "category": "Técnico", "description": "Definición de stack tecnológico", "owner": "Proveedor", "deadline": "Sprint 0"}
            ],
            "dependencies": [
                {"type": "Externa", "description": "APIs de terceros para integraciones"},
                {"type": "Interna", "description": "Aprobación de arquitectura por el cliente"}
            ]
        }
    }


def run_analysis_task(job_id: str, documentation: str, project_name: str, context: str):
    """Tarea en background para análisis y generación de artefactos"""
    try:
        update_job(job_id, status="processing", message="Iniciando análisis...", progress=0.05)
        
        if HAS_ORCHESTRATOR and orchestrator:
            update_job(job_id, message="Procesando con LLM...", progress=0.1)
            
            config = ProjectConfig(
                input_text=documentation,
                project_name=project_name,
                context=context
            )
            
            result = orchestrator.run(config)
            
            update_job(
                job_id,
                status="completed",
                message="Análisis completado exitosamente",
                progress=1.0,
                result=result
            )
        else:
            # Fallback sin LLM
            update_job(job_id, message="Generando artefactos (modo básico)...", progress=0.3)
            result = generate_fallback_result(documentation, project_name)
            
            update_job(
                job_id,
                status="completed",
                message="Análisis completado (modo básico)",
                progress=1.0,
                result=result
            )
            
    except Exception as e:
        logger.error(f"Error en análisis: {e}")
        update_job(
            job_id,
            status="failed",
            message=f"Error: {str(e)}",
            progress=0
        )


# ============================================================
# ENDPOINTS API
# ============================================================

@app.post("/api/analyze", response_model=JobResponse)
async def analyze_project(request: AnalyzeRequest, background_tasks: BackgroundTasks):
    """Inicia el análisis de documentación y generación de artefactos"""
    
    if not request.documentation or len(request.documentation.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="La documentación debe tener al menos 50 caracteres"
        )
    
    job_id = str(uuid.uuid4())
    now = datetime.now().isoformat()
    
    with jobs_lock:
        jobs_store[job_id] = {
            "job_id": job_id,
            "status": "pending",
            "message": "Job creado, iniciando procesamiento...",
            "progress": 0,
            "created_at": now,
            "updated_at": now,
            "result": None
        }
    
    background_tasks.add_task(
        run_analysis_task,
        job_id,
        request.documentation,
        request.project_name or "Proyecto Digital",
        request.context or ""
    )
    
    return JobResponse(
        job_id=job_id,
        status="pending",
        message="Análisis iniciado",
        progress=0,
        created_at=now,
        updated_at=now
    )


@app.get("/api/job/{job_id}", response_model=JobResponse)
async def get_job_status(job_id: str):
    """Obtiene el estado de un job"""
    with jobs_lock:
        if job_id not in jobs_store:
            raise HTTPException(status_code=404, detail="Job no encontrado")
        job = jobs_store[job_id].copy()
    
    return JobResponse(**job)


class PreviewRequest(BaseModel):
    documentation: str = Field(..., description="Contenido del documento para análisis previo")


class PreviewResponse(BaseModel):
    client_name: str
    project_summary: str
    documents_to_generate: List[str]


@app.post("/api/preview", response_model=PreviewResponse)
async def preview_document(request: PreviewRequest):
    """Analiza el documento y devuelve preview con cliente, resumen y documentos a generar"""
    
    if not request.documentation or len(request.documentation.strip()) < 20:
        raise HTTPException(
            status_code=400,
            detail="El documento debe tener contenido suficiente para análisis"
        )
    
    # Lista fija de documentos que se generarán
    documents_list = [
        "🏗️ Diagrama de Arquitectura (Mermaid)",
        "📅 Cronograma del Proyecto (Gantt)",
        "👥 Propuesta de Equipo de Desarrollo",
        "📋 Descomposición en Épicas",
        "💰 Propuesta Económica Estimada",
        "⚠️ Análisis de Riesgos",
        "📝 Prerrequisitos del Proyecto"
    ]
    
    # Intentar usar LLM para extraer cliente y resumen
    try:
        from langchain_openai import AzureChatOpenAI
        
        # Usar credenciales directas de Azure OpenAI (no el proxy AXET)
        llm = AzureChatOpenAI(
            azure_deployment="gpt-4.1",  # Usar el deployment disponible
            openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2025-01-01-preview"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            temperature=0.2,
            max_tokens=800
        )
        
        # Usar más contenido del documento para mejor análisis
        doc_content = request.documentation[:8000]
        
        prompt = f"""Eres un analista experto en proyectos tecnológicos. Tu tarea es analizar cuidadosamente TODA la documentación proporcionada y extraer información PRECISA y UNIFICADA.

CONTEXTO IMPORTANTE:
- La documentación puede contener MÚLTIPLES ARCHIVOS o secciones
- Todos los documentos son APORTES para definir UN SOLO PROYECTO
- Debes SINTETIZAR e INTEGRAR la información de todas las fuentes
- NO analices cada documento por separado, sino como un conjunto que define el proyecto completo

INSTRUCCIONES:
1. Lee TODO el contenido antes de responder
2. Identifica el proyecto ÚNICO que se está definiendo
3. Combina información complementaria de diferentes archivos
4. El resumen debe reflejar la VISIÓN COMPLETA del proyecto, no partes aisladas
5. Si hay información contradictoria, prioriza la más específica o reciente

DOCUMENTACIÓN DEL PROYECTO (puede incluir múltiples archivos):
---
{doc_content}
---

EXTRAE LA SIGUIENTE INFORMACIÓN UNIFICADA:

1. **client_name**: El cliente, empresa u organización del proyecto. Busca en todos los documentos y extrae el nombre más claro. Si no hay mención, indica "Cliente no especificado".

2. **project_summary**: Un resumen ejecutivo CONSOLIDADO de 4-6 oraciones que:
   - Describa el proyecto como UN TODO coherente
   - Integre los objetivos y alcance mencionados en los diferentes documentos
   - Destaque las funcionalidades PRINCIPALES del sistema
   - Mencione tecnologías o integraciones relevantes
   - Sea ESPECÍFICO a este proyecto (evita frases genéricas)

IMPORTANTE: Tu resumen debe leer como si fuera la descripción de UN proyecto, NO como un listado de lo que dice cada documento.

Responde ÚNICAMENTE con un JSON válido:
{{"client_name": "nombre del cliente", "project_summary": "resumen UNIFICADO del proyecto"}}
"""
        
        response = llm.invoke(prompt)
        content = response.content.strip()
        
        # Extraer JSON de la respuesta (manejar JSONs multilínea)
        import re
        json_match = re.search(r'\{[\s\S]*?"client_name"[\s\S]*?"project_summary"[\s\S]*?\}', content)
        if json_match:
            json_str = json_match.group()
            json_str = json_str.replace('\n', ' ').replace('\r', ' ')
            data = json.loads(json_str)
            
            client = data.get("client_name", "").strip()
            summary = data.get("project_summary", "").strip()
            
            if not client or client.lower() in ["", "n/a", "no especificado"]:
                client = "Cliente no especificado"
            if not summary or len(summary) < 20:
                summary = "Proyecto para análisis de artefactos. Revisar documento fuente para más detalles."
            
            return PreviewResponse(
                client_name=client,
                project_summary=summary,
                documents_to_generate=documents_list
            )
    except Exception as e:
        logger.warning(f"Error en preview LLM: {e}")
    
    # Fallback: análisis básico sin LLM
    doc_lower = request.documentation.lower()
    
    # Intentar detectar nombre de cliente
    client_name = "Cliente no especificado"
    for keyword in ["cliente:", "empresa:", "organización:", "compañía:", "client:"]:
        if keyword in doc_lower:
            idx = doc_lower.index(keyword) + len(keyword)
            end = doc_lower.find("\n", idx)
            if end == -1:
                end = idx + 50
            client_name = request.documentation[idx:end].strip()[:50]
            break
    
    # Generar resumen básico
    first_lines = request.documentation[:500].replace("\n", " ").strip()
    project_summary = first_lines[:200] + "..." if len(first_lines) > 200 else first_lines
    
    return PreviewResponse(
        client_name=client_name,
        project_summary=project_summary,
        documents_to_generate=documents_list
    )


# ============================================================
# ARTIFACT ENDPOINTS - Individual artifact generation
# ============================================================

class ArtifactRequest(BaseModel):
    documentation: str
    project_name: str = "Proyecto Digital"
    additional_context: Optional[str] = None  # Contexto adicional del usuario para los LLMs

class ArchitectureResponse(BaseModel):
    artifact_type: str = "architecture"
    mermaid_code: str
    description: str
    layers: Optional[dict] = None  # Resumen de componentes por capa


@app.post("/api/artifact/architecture", response_model=ArchitectureResponse)
async def generate_architecture(request: ArtifactRequest):
    """Genera un diagrama de arquitectura C4 en formato Mermaid con capas arquitectónicas"""
    
    if not request.documentation or len(request.documentation.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="La documentación debe tener contenido suficiente"
        )
    
    try:
        from langchain_openai import AzureChatOpenAI
        
        llm = AzureChatOpenAI(
            azure_deployment="gpt-4.1",
            openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2025-01-01-preview"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            temperature=0.3,
            max_tokens=3500
        )
        
        doc_content = request.documentation[:6000]
        
        # Construir sección de contexto adicional si existe
        context_section = ""
        if request.additional_context and request.additional_context.strip():
            context_section = f"""
=== CONTEXTO ADICIONAL DEL USUARIO ===
{request.additional_context.strip()}

IMPORTANTE: Considera este contexto adicional al generar el diagrama de arquitectura.
"""
        
        prompt = f"""Eres un Arquitecto de Software Senior. Genera un diagrama de arquitectura C4 (nivel Container) en Mermaid.

DOCUMENTACIÓN DEL PROYECTO:
{doc_content}
{context_section}

=== CAPAS ARQUITECTÓNICAS ESTÁNDAR (OBLIGATORIO incluir las que apliquen) ===

1. 🖥️ CAPA DE PRESENTACIÓN / CONSUMIDORES
   - Aplicaciones Web (SPA, MPA)
   - Aplicaciones Móviles (iOS, Android)
   - Portales, PWA
   - Clientes API externos

2. 🔐 CAPA DE SEGURIDAD (IAM)
   - Identity Provider (Azure AD, Keycloak, Auth0)
   - API Gateway con autenticación
   - SSO, OAuth 2.0, OIDC
   - Gestión de permisos y roles

3. ⚙️ CAPA BACKEND / SERVICIOS
   - APIs REST / GraphQL
   - Microservicios
   - Servicios de dominio
   - Lógica de negocio

4. 🗄️ CAPA DE ENTORNO DE DATOS
   - Bases de datos relacionales (SQL Server, PostgreSQL)
   - Bases de datos NoSQL (MongoDB, CosmosDB)
   - Cache (Redis)
   - Data Lakes, Storage

5. 🔗 CAPA DE INTEGRACIONES
   - APIs de terceros
   - Servicios internos (legacy)
   - Message Brokers (Kafka, Service Bus)
   - ETL, ESB

6. 📊 CAPA DE REPORTES / ANALYTICS
   - Dashboards (Power BI, Grafana)
   - Data Warehouse
   - Servicios de BI
   - Logs y Monitoreo

=== FORMATO MERMAID C4 REQUERIDO ===

Usa el formato C4Context de Mermaid con Boundary para cada capa:

```
C4Context
    title Arquitectura del Sistema - [Nombre Proyecto]
    
    Enterprise_Boundary(b0, "Sistema") {{
        
        Boundary(presentation, "Capa de Presentación", "🖥️") {{
            Person(user, "Usuario", "Descripción")
            System(webapp, "Aplicación Web", "React/Angular")
        }}
        
        Boundary(security, "Capa de Seguridad", "🔐") {{
            System(idp, "Identity Provider", "Azure AD")
            System(apigw, "API Gateway", "Kong/APIM")
        }}
        
        Boundary(backend, "Capa Backend", "⚙️") {{
            System(api, "API Principal", "FastAPI/Spring")
            System(svc, "Servicio Core", "Microservicio")
        }}
        
        Boundary(data, "Capa de Datos", "🗄️") {{
            SystemDb(db, "Base de Datos", "PostgreSQL")
            SystemDb(cache, "Cache", "Redis")
        }}
        
        Boundary(integrations, "Capa Integraciones", "🔗") {{
            System_Ext(ext1, "API Externa", "Tercero")
            System(broker, "Message Broker", "Kafka")
        }}
        
        Boundary(reports, "Capa Reportes", "📊") {{
            System(bi, "Dashboard", "Power BI")
        }}
    }}
    
    Rel(user, webapp, "Usa")
    Rel(webapp, apigw, "Consume API")
    Rel(apigw, idp, "Valida token")
    Rel(apigw, api, "Enruta")
    Rel(api, db, "Lee/Escribe")
```

=== RESPUESTA JSON ===

{{
    "mermaid_code": "C4Context\\n    title Arquitectura...\\n    ...",
    "description": "Descripción de la arquitectura identificada",
    "layers": {{
        "presentation": ["Componente1", "Componente2"],
        "security": ["Azure AD", "API Gateway"],
        "backend": ["API Core", "Servicio X"],
        "data": ["PostgreSQL", "Redis"],
        "integrations": ["API Externa Y"],
        "reports": ["Power BI Dashboard"]
    }}
}}

IMPORTANTE: 
- Responde SOLO con JSON válido (sin texto antes ni después)
- NO incluyas bloques de código markdown
- Incluye SOLO las capas que identifiques en el proyecto
- Los IDs de componentes deben ser simples (sin espacios ni caracteres especiales)
- Usa System_Ext para sistemas externos
- Usa SystemDb para bases de datos"""
        
        response = llm.invoke(prompt)
        content = response.content.strip()
        
        # Limpiar respuesta de posibles bloques de código
        if content.startswith("```"):
            lines = content.split("\n")
            content = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
            content = content.strip()
        
        # Buscar JSON con llaves balanceadas
        def extract_json(text):
            """Extrae el primer objeto JSON válido del texto"""
            start = text.find("{")
            if start == -1:
                return None
            
            depth = 0
            for i, char in enumerate(text[start:], start):
                if char == "{":
                    depth += 1
                elif char == "}":
                    depth -= 1
                    if depth == 0:
                        return text[start:i+1]
            return None
        
        json_str = extract_json(content)
        
        if json_str:
            try:
                data = json.loads(json_str)
                
                mermaid = data.get("mermaid_code", "").strip()
                desc = data.get("description", "").strip()
                layers = data.get("layers", {})
                
                if mermaid:
                    return ArchitectureResponse(
                        mermaid_code=mermaid,
                        description=desc or "Diagrama de arquitectura C4 generado automáticamente",
                        layers=layers
                    )
            except json.JSONDecodeError as je:
                logger.error(f"Error parsing JSON: {je}")
                logger.error(f"JSON string: {json_str[:500]}...")
        
        # Fallback si no se puede parsear
        raise HTTPException(status_code=500, detail="Error generando diagrama: respuesta inválida del LLM")
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generando arquitectura: {e}")
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


# ============================================================
# TEAM PROPOSAL ENDPOINT
# ============================================================

AVAILABLE_PROFILES = """
- Administrador Middleware
- Administrador Openshift / Kubernetes
- Administrador plataforma CRM Dynamics
- Agile Coach
- Analista de Negocio Junior/Senior
- Arquitecto / Experto DevOps
- Arquitecto de contenedorizacion (Docker, Kubernetes, Microservicios)
- Arquitecto de Integracion (ESB, Queues, Serverless)
- Arquitecto especialista (Azure, AWS, RedHat, Oracle, IBM, Google Cloud)
- Automatizador de pruebas funcionales Jr/Sr
- Business Intelligence developer
- CloudOps Admin
- CRM Dynamics developer junior/senior
- Data Engineer Advance Analytics
- Data Engineer Analytics Azure (Data Bricks for Analytics)
- Data Engineer Analytics OCI (Oracle cloud Infraestructure)
- Data Engineer Azure Junior/Senior (ADF, SQL, Data Lake, Databricks)
- Data Engineer Oracle Junior/Senior (SQL y PLSQL)
- Data Governance Specialist
- Data Management Specialist
- Data Scientist ML - IA
- DBA Oracle
- DBA PostgreSQL
- DBA Sql Server
- Desarrollador Umbraco junior/senior
- Desarrollador .Net Junior/Senior
- Desarrollador Android Junior/Senior
- Desarrollador Angular Junior/Senior
- Desarrollador Aplicaciones Moviles Junior/Senior
- Desarrollador Back AEM junior/senior
- Desarrollador Base de Datos Junior/Senior
- Desarrollador Front AEM junior/senior
- Desarrollador Front Junior/Senior
- Desarrollador full stack Azure Junior/Senior
- Desarrollador IOS Junior/Senior
- Desarrollador Java Junior/Senior
- Desarrollador JBOSS Fuse Junior/Senior
- Desarrollador Power BI
- Desarrollador PowerShell / YAML / CLI en Linux
- Desarrollador PWA Junior/Senior
- Desarrollador SharePoint
- Desarrollador Spring Boot / Quarkus / Apache camel Junior/Senior
- Desarrollador Xamarin Junior/Senior
- Disenador UX Junior/Senior
- Experto DevOps
- Experto tecnico de plataforma AEM para desarrollo (Campaing, target, etc)
- Gestor de Proyecto Junior/Senior
- Machine Learning and Architect Engineer
- Machine Learning Engineer
- Project Management
- Scrum Master
- Technical Developer Lead
- Tester de pruebas funcionales Jr/Sr
- UI Designer
"""

class TeamMember(BaseModel):
    role: str
    seniority: str
    quantity: int
    justification: str

class TeamResponse(BaseModel):
    artifact_type: str = "team"
    team_members: List[dict]
    total_members: int
    summary: str


@app.post("/api/artifact/team", response_model=TeamResponse)
async def generate_team_proposal(request: ArtifactRequest):
    """Genera una propuesta de equipo basada en la documentación"""
    
    if not request.documentation or len(request.documentation.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="La documentación debe tener contenido suficiente"
        )
    
    try:
        from langchain_openai import AzureChatOpenAI
        
        llm = AzureChatOpenAI(
            azure_deployment="gpt-4.1",
            openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2025-01-01-preview"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            temperature=0.3,
            max_tokens=3000
        )
        
        doc_content = request.documentation[:6000]
        
        prompt = f"""Eres un experto en gestión de proyectos de software. Analiza la documentación del proyecto y sugiere el equipo necesario.

DOCUMENTACIÓN DEL PROYECTO:
{doc_content}

PERFILES DISPONIBLES (SOLO puedes elegir de esta lista):
{AVAILABLE_PROFILES}

INSTRUCCIONES:
1. Analiza las tecnologías, alcance y complejidad del proyecto
2. Selecciona SOLO perfiles de la lista proporcionada
3. Indica nivel de seniority (Junior o Senior)
4. Indica cantidad de personas por perfil
5. Justifica brevemente por qué se necesita cada perfil

Responde SOLO con JSON válido:
{{
    "team_members": [
        {{"role": "Nombre exacto del perfil", "seniority": "Junior/Senior", "quantity": 1, "justification": "Razón breve"}},
        ...
    ],
    "summary": "Resumen de 2-3 oraciones sobre la composición del equipo sugerido"
}}"""
        
        response = llm.invoke(prompt)
        content = response.content.strip()
        
        # Extraer JSON
        import re
        json_match = re.search(r'\{[\s\S]*?"team_members"[\s\S]*?\][\s\S]*?"summary"[\s\S]*?\}', content)
        
        if json_match:
            json_str = json_match.group()
            data = json.loads(json_str)
            
            team = data.get("team_members", [])
            summary = data.get("summary", "")
            total = sum(m.get("quantity", 1) for m in team)
            
            if team:
                return TeamResponse(
                    team_members=team,
                    total_members=total,
                    summary=summary or f"Equipo propuesto de {total} miembros"
                )
        
        raise HTTPException(status_code=500, detail="Error generando propuesta de equipo")
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generando equipo: {e}")
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


# ============================================================
# EPICS ENDPOINT
# ============================================================

class EpicsResponse(BaseModel):
    artifact_type: str = "epics"
    markdown_content: str
    epics_count: int
    summary: str


@app.post("/api/artifact/epics", response_model=EpicsResponse)
async def generate_epics(request: ArtifactRequest):
    """Genera las épicas del proyecto en formato Markdown"""
    
    if not request.documentation or len(request.documentation.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="La documentación debe tener contenido suficiente"
        )
    
    try:
        from langchain_openai import AzureChatOpenAI
        
        llm = AzureChatOpenAI(
            azure_deployment="gpt-4.1",
            openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2025-01-01-preview"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            temperature=0.3,
            max_tokens=4000
        )
        
        doc_content = request.documentation[:8000]
        
        prompt = f"""Eres un Product Owner experto. Analiza la documentación del proyecto y genera una descomposición completa en Épicas.

DOCUMENTACIÓN DEL PROYECTO:
{doc_content}

INSTRUCCIONES:
1. Identifica todas las funcionalidades principales del proyecto
2. Agrupa en épicas lógicas (máximo 10-15 épicas)
3. Para cada épica incluye:
   - Título claro
   - Descripción del objetivo
   - Historias de usuario principales (3-5 por épica)
   - Criterios de aceptación generales
   - Dependencias con otras épicas (si aplica)

FORMATO MARKDOWN REQUERIDO:
Genera el contenido en Markdown limpio siguiendo este formato:

# Épicas del Proyecto: [Nombre del Proyecto]

## Resumen Ejecutivo
[Breve descripción del alcance]

---

## Épica 1: [Nombre de la Épica]

### Descripción
[Objetivo de la épica]

### Historias de Usuario
- **HU-1.1**: Como [usuario], quiero [funcionalidad] para [beneficio]
- **HU-1.2**: ...

### Criterios de Aceptación
- [ ] Criterio 1
- [ ] Criterio 2

### Dependencias
- Épica X (si aplica)

---

## Épica 2: ...

[continúa con todas las épicas]

Responde SOLO con el contenido Markdown (sin bloques de código ni JSON)."""
        
        response = llm.invoke(prompt)
        markdown_content = response.content.strip()
        
        # Limpiar el markdown si viene con bloques de código
        if markdown_content.startswith("```"):
            lines = markdown_content.split("\n")
            markdown_content = "\n".join(lines[1:-1] if lines[-1] == "```" else lines[1:])
        
        # Contar épicas
        epics_count = markdown_content.count("## Épica")
        if epics_count == 0:
            epics_count = markdown_content.count("## Epic")
        
        return EpicsResponse(
            markdown_content=markdown_content,
            epics_count=epics_count,
            summary=f"Se identificaron {epics_count} épicas principales para el proyecto"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generando épicas: {e}")
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


class DownloadEpicsRequest(BaseModel):
    markdown_content: str
    project_name: str = "Proyecto"


@app.post("/api/artifact/epics/download")
async def download_epics_word(request: DownloadEpicsRequest):
    """Convierte el Markdown de épicas a documento Word"""
    import io
    
    if not DOCX_AVAILABLE:
        logger.error("Módulo python-docx no instalado")
        raise HTTPException(status_code=500, detail="Módulo python-docx no instalado. Instalar con: pip install python-docx")
    
    try:
        doc = Document()
        
        # Título del documento
        title = doc.add_heading('Épicas del Proyecto', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Procesar Markdown línea por línea
        lines = request.markdown_content.split('\n')
        
        for line in lines:
            try:
                line = line.strip()
                
                if not line or line == '---':
                    continue
                    
                # Encabezados
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                # Listas con checkbox
                elif line.startswith('- [ ]') or line.startswith('- [x]'):
                    text = line[5:].strip()
                    prefix = '☐ ' if '[ ]' in line else '☑ '
                    add_markdown_paragraph(doc, prefix + text, style='List Bullet')
                # Listas normales
                elif line.startswith('- '):
                    add_markdown_paragraph(doc, line[2:], style='List Bullet')
                # Párrafos normales
                else:
                    if line.strip():
                        add_markdown_paragraph(doc, line)
            except Exception as line_error:
                logger.warning(f"Error procesando línea: {line_error}")
                continue
        
        # Guardar en memoria
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Nombre de archivo seguro
        safe_name = request.project_name.replace(' ', '_').replace('/', '_')
        
        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=epicas_{safe_name}.docx"
            }
        )
        
    except Exception as e:
        logger.error(f"Error generando Word: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")



# ============================================================
# GANTT CHRONOGRAM ENDPOINT
# ============================================================

# Capas arquitectónicas estándar para el cronograma
ARCHITECTURAL_LAYERS = {
    "presentation": {
        "name": "Capa de Presentación / Consumidores",
        "icon": "🖥️",
        "color": "#4A90D9",
        "description": "UI, Apps móviles, Web, Portales, PWA"
    },
    "security": {
        "name": "Capa de Seguridad (IAM)",
        "icon": "🔐",
        "color": "#E74C3C",
        "description": "Autenticación, Autorización, SSO, OAuth, Gestión de identidades"
    },
    "backend": {
        "name": "Capa Backend",
        "icon": "⚙️",
        "color": "#50C878",
        "description": "Servicios, APIs REST/GraphQL, Microservicios, Lógica de negocio"
    },
    "data": {
        "name": "Capa de Entorno de Datos",
        "icon": "🗄️",
        "color": "#9B59B6",
        "description": "Bases de datos, Data Lakes, Cache, Modelado de datos"
    },
    "integrations": {
        "name": "Capa de Integraciones",
        "icon": "🔗",
        "color": "#F5A623",
        "description": "APIs externas, Servicios internos, ETL, Messaging, ESB"
    },
    "reports": {
        "name": "Capa de Reportes",
        "icon": "📊",
        "color": "#1ABC9C",
        "description": "Dashboards, BI, Analytics, Visualizaciones, KPIs"
    }
}

class GanttRequest(BaseModel):
    documentation: str
    project_name: str = "Proyecto"
    start_date: Optional[str] = None  # Fecha de inicio del proyecto (YYYY-MM-DD)
    epics_data: Optional[dict] = None
    team_data: Optional[dict] = None
    architecture_data: Optional[dict] = None
    additional_context: Optional[str] = None  # Contexto adicional del usuario

class GanttTask(BaseModel):
    id: str
    name: str
    start_date: str
    end_date: str
    duration_days: int
    epic: str
    layer: str  # NUEVO: Capa arquitectónica

class GanttResponse(BaseModel):
    artifact_type: str = "gantt"
    mermaid_code: str
    tasks: List[dict]
    total_weeks: int
    summary: str
    layers_summary: Optional[dict] = None  # NUEVO: Resumen por capas
    architectural_layers: Optional[dict] = None  # NUEVO: Definición de capas


@app.post("/api/artifact/gantt", response_model=GanttResponse)
async def generate_gantt(request: GanttRequest):
    """Genera un cronograma Gantt basado en épicas, equipo y arquitectura"""
    
    if not request.documentation or len(request.documentation.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="La documentación debe tener contenido suficiente"
        )
    
    try:
        from langchain_openai import AzureChatOpenAI
        
        llm = AzureChatOpenAI(
            azure_deployment="gpt-4.1",
            openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2025-01-01-preview"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            temperature=0.3,
            max_tokens=6000  # Aumentado para evitar truncamiento del JSON
        )
        
        # Construir contexto con artefactos previos
        context_parts = [f"DOCUMENTACIÓN:\n{request.documentation[:4000]}"]
        
        if request.epics_data:
            epics_summary = request.epics_data.get("markdown_content", "")[:2000]
            context_parts.append(f"\nÉPICAS IDENTIFICADAS:\n{epics_summary}")
        
        if request.team_data:
            team_members = request.team_data.get("team_members", [])
            team_summary = ", ".join([f"{m.get('role')} ({m.get('quantity')})" for m in team_members[:10]])
            context_parts.append(f"\nEQUIPO PROPUESTO: {team_summary}")
        
        # Agregar contexto adicional del usuario si existe
        if request.additional_context and request.additional_context.strip():
            context_parts.append(f"\n=== CONTEXTO ADICIONAL DEL USUARIO ===\n{request.additional_context.strip()}\n\nIMPORTANTE: Considera este contexto al generar el cronograma.")
        
        context = "\n".join(context_parts)
        
        # Determinar fecha de inicio
        start_date = request.start_date or datetime.now().strftime("%Y-%m-%d")
        
        # Construir descripción de capas para el prompt
        layers_description = "\n".join([
            f"- {key}: {info['icon']} {info['name']} - {info['description']}"
            for key, info in ARCHITECTURAL_LAYERS.items()
        ])
        
        prompt = f"""Eres un Project Manager y Arquitecto de Software experto. Genera un cronograma detallado organizado por CAPAS ARQUITECTÓNICAS.

{context}

FECHA DE INICIO DEL PROYECTO: {start_date}

=== CAPAS ARQUITECTÓNICAS (OBLIGATORIO usar estas) ===
{layers_description}

=== INSTRUCCIONES ===
1. CADA tarea debe estar asignada a UNA de las 6 capas arquitectónicas (campo "layer")
2. Agrupa las tareas primero por CAPA, luego por épica dentro de cada capa
3. Cada tarea debe tener un PERFIL asignado (ej: "Arquitecto", "Desarrollador Senior")
4. Calcula fechas REALES basadas en la fecha de inicio: {start_date}
5. Considera dependencias entre capas (ej: Backend antes de Presentación)
6. El código Mermaid debe usar "section" para cada capa arquitectónica

=== ESTRUCTURA REQUERIDA para cada tarea ===
- id: Identificador único (task1, task2, etc.)
- layer: Capa arquitectónica (presentation, security, backend, data, integrations, reports)
- general_epic: Nombre de la fase/épica general
- specific_epic: Componente específico dentro de la capa
- name: Nombre descriptivo de la tarea
- profile: Perfil/rol que ejecuta la tarea
- start_date: Fecha de inicio (YYYY-MM-DD)
- end_date: Fecha de fin (YYYY-MM-DD)
- duration_days: Duración en días laborales

=== RESPUESTA JSON (OBLIGATORIO) ===
Responde ÚNICAMENTE con un JSON válido sin bloques de código markdown. El JSON debe tener esta estructura exacta:

{{"mermaid_code": "gantt\\n    title Cronograma\\n    dateFormat YYYY-MM-DD\\n    section Capa1\\n    Tarea1 :t1, {start_date}, 5d", "tasks": [{{"id": "t1", "layer": "backend", "general_epic": "Desarrollo", "specific_epic": "APIs", "name": "Crear APIs", "profile": "Desarrollador", "start_date": "{start_date}", "end_date": "2025-02-07", "duration_days": 5, "epic": "Desarrollo"}}], "total_weeks": 12, "summary": "Cronograma de 12 semanas"}}

REGLAS ESTRICTAS:
1. NO uses bloques de código (```) 
2. Responde SOLO con el JSON, sin texto antes ni después
3. El campo "layer" DEBE ser uno de: presentation, security, backend, data, integrations, reports
4. Incluye tareas para las capas que apliquen al proyecto"""
        
        response = llm.invoke(prompt)
        content = response.content.strip()
        
        # LOG: Ver qué devuelve el LLM
        logger.info(f"=== GANTT LLM Response (primeros 500 chars) ===")
        logger.info(content[:500])
        
        # Limpiar respuesta de posibles bloques de código markdown
        if "```json" in content:
            # Extraer contenido entre ```json y ```
            import re
            json_block = re.search(r'```json\s*([\s\S]*?)\s*```', content)
            if json_block:
                content = json_block.group(1).strip()
        elif "```" in content:
            # Remover cualquier bloque de código
            lines = content.split("\n")
            new_lines = []
            in_code_block = False
            for line in lines:
                if line.strip().startswith("```"):
                    in_code_block = not in_code_block
                elif not in_code_block:
                    new_lines.append(line)
            content = "\n".join(new_lines).strip()
        
        # Si empieza con ``` sin json, limpiar
        if content.startswith("```"):
            lines = content.split("\n")
            content = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
            content = content.strip()
        
        # Buscar JSON con llaves balanceadas
        def extract_json(text):
            """Extrae el primer objeto JSON válido del texto"""
            start = text.find("{")
            if start == -1:
                return None
            
            depth = 0
            in_string = False
            escape_next = False
            
            for i, char in enumerate(text[start:], start):
                if escape_next:
                    escape_next = False
                    continue
                    
                if char == '\\':
                    escape_next = True
                    continue
                    
                if char == '"' and not escape_next:
                    in_string = not in_string
                    continue
                
                if not in_string:
                    if char == "{":
                        depth += 1
                    elif char == "}":
                        depth -= 1
                        if depth == 0:
                            return text[start:i+1]
            return None
        
        # Si el contenido ya empieza con {, intentar parse directo primero
        json_str = None
        if content.strip().startswith("{"):
            json_str = content.strip()
        else:
            json_str = extract_json(content)
        
        logger.info(f"=== Extracted JSON (primeros 300 chars) ===")
        logger.info(json_str[:300] if json_str else "NO JSON FOUND")
        
        if json_str:
            try:
                data = json.loads(json_str)
                
                mermaid = data.get("mermaid_code", "").strip()
                tasks = data.get("tasks", [])
                total_weeks = data.get("total_weeks", 0)
                summary = data.get("summary", "")
                
                # Calcular resumen por capas
                layers_summary = {}
                for task in tasks:
                    layer = task.get("layer", "backend")
                    if layer not in layers_summary:
                        layers_summary[layer] = {
                            "count": 0,
                            "total_days": 0,
                            "name": ARCHITECTURAL_LAYERS.get(layer, {}).get("name", layer),
                            "icon": ARCHITECTURAL_LAYERS.get(layer, {}).get("icon", "📦"),
                            "color": ARCHITECTURAL_LAYERS.get(layer, {}).get("color", "#666666")
                        }
                    layers_summary[layer]["count"] += 1
                    layers_summary[layer]["total_days"] += task.get("duration_days", 0)
                
                if mermaid:
                    return GanttResponse(
                        mermaid_code=mermaid,
                        tasks=tasks,
                        total_weeks=total_weeks,
                        summary=summary or f"Cronograma de {total_weeks} semanas organizado por capas arquitectónicas",
                        layers_summary=layers_summary,
                        architectural_layers=ARCHITECTURAL_LAYERS
                    )
            except json.JSONDecodeError as je:
                logger.error(f"Error parsing Gantt JSON: {je}")
                logger.error(f"JSON string: {json_str[:500]}...")
        
        raise HTTPException(status_code=500, detail="Error generando cronograma: respuesta inválida del LLM")
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generando Gantt: {e}")
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


class DownloadGanttRequest(BaseModel):
    tasks: List[dict]
    project_name: str = "Proyecto"


@app.post("/api/artifact/gantt/download")
async def download_gantt_excel(request: DownloadGanttRequest):
    """Exporta el cronograma Gantt a Excel con timeline visual"""
    import io
    from datetime import datetime, timedelta
    
    if not XLSX_AVAILABLE:
        raise HTTPException(status_code=500, detail="openpyxl no instalado")
    
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Cronograma"
        
        # Colores para épicas
        EPIC_COLORS = [
            "4A90D9",  # Azul
            "50C878",  # Verde
            "F5A623",  # Naranja
            "9B59B6",  # Púrpura
            "E74C3C",  # Rojo
            "1ABC9C",  # Turquesa
        ]
        
        # Estilos
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="090f26", end_color="090f26", fill_type="solid")
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Calcular rango de semanas
        all_dates = []
        for task in request.tasks:
            try:
                start = datetime.strptime(task.get("start_date", ""), "%Y-%m-%d")
                end = datetime.strptime(task.get("end_date", ""), "%Y-%m-%d")
                all_dates.extend([start, end])
            except:
                pass
        
        if all_dates:
            project_start = min(all_dates)
            project_end = max(all_dates)
            total_weeks = max(1, ((project_end - project_start).days // 7) + 1)
        else:
            project_start = datetime.now()
            total_weeks = 12
        
        # Headers: columnas fijas + semanas
        fixed_headers = ["Épica General", "Épica Específica", "Tarea", "Perfil", "Inicio", "Fin", "Días"]
        week_headers = [f"S{i+1}" for i in range(min(total_weeks, 24))]  # Max 24 semanas
        all_headers = fixed_headers + week_headers
        
        # Escribir headers
        for col, header in enumerate(all_headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border
        
        # Mapear épicas a colores
        epic_color_map = {}
        color_index = 0
        
        # Escribir datos y timeline
        for row, task in enumerate(request.tasks, 2):
            # Datos básicos
            general_epic = task.get("general_epic", task.get("epic", ""))
            specific_epic = task.get("specific_epic", "")
            
            # Asignar color a épica general
            if general_epic not in epic_color_map:
                epic_color_map[general_epic] = EPIC_COLORS[color_index % len(EPIC_COLORS)]
                color_index += 1
            
            epic_color = epic_color_map[general_epic]
            epic_fill = PatternFill(start_color=epic_color, end_color=epic_color, fill_type="solid")
            
            ws.cell(row=row, column=1, value=general_epic).border = border
            ws.cell(row=row, column=2, value=specific_epic).border = border
            ws.cell(row=row, column=3, value=task.get("name", "")).border = border
            ws.cell(row=row, column=4, value=task.get("profile", "")).border = border
            ws.cell(row=row, column=5, value=task.get("start_date", "")).border = border
            ws.cell(row=row, column=6, value=task.get("end_date", "")).border = border
            ws.cell(row=row, column=7, value=task.get("duration_days", 0)).border = border
            
            # Timeline visual - colorear celdas de semanas donde hay actividad
            try:
                task_start = datetime.strptime(task.get("start_date", ""), "%Y-%m-%d")
                task_end = datetime.strptime(task.get("end_date", ""), "%Y-%m-%d")
                
                for week_idx in range(min(total_weeks, 24)):
                    week_start = project_start + timedelta(weeks=week_idx)
                    week_end = week_start + timedelta(days=6)
                    
                    col = 8 + week_idx  # Columna de la semana
                    cell = ws.cell(row=row, column=col)
                    cell.border = border
                    
                    # Si la tarea está activa en esta semana, colorear
                    if task_start <= week_end and task_end >= week_start:
                        cell.fill = epic_fill
                        cell.value = "●"
                        cell.alignment = Alignment(horizontal="center")
            except:
                pass
        
        # Ajustar anchos de columnas
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 35
        ws.column_dimensions['D'].width = 25
        ws.column_dimensions['E'].width = 12
        ws.column_dimensions['F'].width = 12
        ws.column_dimensions['G'].width = 6
        
        # Semanas más estrechas
        for i in range(min(total_weeks, 24)):
            col_letter = get_column_letter(8 + i)
            ws.column_dimensions[col_letter].width = 4
        
        # Guardar
        file_stream = io.BytesIO()
        wb.save(file_stream)
        file_stream.seek(0)
        
        safe_name = request.project_name.replace(' ', '_').replace('/', '_')
        
        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={
                "Content-Disposition": f"attachment; filename=cronograma_{safe_name}.xlsx"
            }
        )
        
    except Exception as e:
        logger.error(f"Error generando Excel: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


# ============================================================
# ECONOMICS ENDPOINT
# ============================================================

# Tabla de tarifas por día en COP
DAILY_RATES_COP = {
    "administrador middleware": 123883,
    "administrador openshift": 123883,
    "administrador openshift / kubernetes": 123883,
    "administrador plataforma crm dynamics": 148359,
    "agile coach": 91238,
    "analista de negocio junior": 49024,
    "analista de negocio senior": 68273,
    "arquitecto / experto devops": 144195,
    "arquitecto de contenedorizacion": 149086,
    "arquitecto de integracion": 139307,
    "arquitecto especialista": 180689,
    "automatizador de pruebas funcionales": 61674,
    "automatizador pruebas no funcional jr": 56348,
    "automatizador pruebas no funcional sr": 79785,
    "business intelligence developer": 74475,
    "cloudops admin": 94149,
    "crm dynamics developer junior": 76508,
    "crm dynamics developer senior": 124638,
    "data engineer advance analytics": 158601,
    "data engineer analytics azure": 142215,
    "data engineer analytics oci": 143840,
    "data engineer azure junior": 75510,
    "data engineer azure senior": 147856,
    "data engineer oracle junior": 75510,
    "data engineer oracle senior": 93857,
    "data governance specialist": 158277,
    "data management specialist": 128447,
    "data scientist ml - ia": 130740,
    "dba oracle": 75510,
    "dba postgresql": 75510,
    "dba sql server": 75510,
    "desarrollador umbraco senior": 110011,
    "desarrollador .net junior": 50489,
    "desarrollador .net senior": 86378,
    "desarrollador android junior": 50489,
    "desarrollador android senior": 95298,
    "desarrollador angular junior": 51223,
    "desarrollador angular senior": 92746,
    "desarrollador aplicaciones moviles junior": 76508,
    "desarrollador aplicaciones moviles senior": 101820,
    "desarrollador back aem junior": 80808,
    "desarrollador back aem senior": 154827,
    "desarrollador base de datos junior": 62092,
    "desarrollador base de datos senior": 90792,
    "desarrollador front aem junior": 71256,
    "desarrollador front aem senior": 145362,
    "desarrollador front junior": 60803,
    "desarrollador front senior": 97627,
    "desarrollador full stack azure junior": 67800,
    "desarrollador full stack azure senior": 100456,
    "desarrollador ios junior": 50489,
    "desarrollador ios senior": 95298,
    "desarrollador java junior": 49025,
    "desarrollador java senior": 92746,
    "desarrollador jboss fuse junior": 63307,
    "desarrollador jboss fuse senior": 100294,
    "desarrollador power bi": 72906,
    "desarrollador powershell / yaml / cli en linux": 103443,
    "desarrollador pwa junior": 67800,
    "desarrollador pwa senior": 113229,
    "desarrollador sharepoint": 92746,
    "desarrollador spring boot junior": 55756,
    "desarrollador spring boot senior": 89580,
    "desarrollador spring boot / quarkus junior": 55756,
    "desarrollador spring boot / quarkus senior": 89580,
    "desarrollador umbraco junior": 76886,
    "desarrollador xamarin junior": 71256,
    "desarrollador xamarin senior": 113229,
    "disenador ux junior": 73748,
    "disenador ux senior": 96930,
    "diseñador ux junior": 73748,
    "diseñador ux senior": 96930,
    "experto devops": 157237,
    "experto tecnico de plataforma aem": 217489,
    "gestor de proyecto junior": 82512,
    "gestor de proyecto senior": 110530,
    "machine learning and architect engineer": 187304,
    "machine learning engineer": 104994,
    "project management": 139341,
    "scrum master": 88519,
    "technical developer lead": 96819,
    "tester de pruebas funcionales jr": 35282,
    "tester de pruebas funcionales sr": 53419,
    "ui designer": 80808,
}

class EconomicsRequest(BaseModel):
    documentation: str
    project_name: str = "Proyecto"
    team_data: Optional[dict] = None
    gantt_data: Optional[dict] = None

class EconomicsResponse(BaseModel):
    artifact_type: str = "economics"
    line_items: List[dict]
    total_cost: float
    total_days: int
    currency: str = "COP"
    summary: str


@app.post("/api/artifact/economics", response_model=EconomicsResponse)
async def generate_economics(request: EconomicsRequest):
    """Genera propuesta económica alineada con el cronograma Gantt"""
    
    line_items = []
    total_cost = 0
    total_days = 0
    
    # NUEVO: Calcular días por perfil desde el Gantt
    profile_days = {}  # {profile: total_days}
    
    if request.gantt_data and request.gantt_data.get("tasks"):
        for task in request.gantt_data["tasks"]:
            profile = task.get("profile", "").strip()
            duration = task.get("duration_days", 0)
            
            if profile and duration > 0:
                # Normalizar nombre del perfil para matching
                profile_key = profile.lower().strip()
                if profile_key not in profile_days:
                    profile_days[profile_key] = {"name": profile, "days": 0}
                profile_days[profile_key]["days"] += duration
    
    # Si tenemos perfiles del Gantt, usarlos
    if profile_days:
        for profile_key, profile_info in profile_days.items():
            profile_name = profile_info["name"]
            days = profile_info["days"]
            
            # Buscar tarifa por perfil
            daily_rate = None
            profile_lower = profile_name.lower()
            
            # Búsqueda en tabla de tarifas
            for rate_role, rate in DAILY_RATES_COP.items():
                # Coincidir palabras clave
                if any(word in rate_role for word in profile_lower.split()[:3]):
                    daily_rate = rate
                    break
                if any(word in profile_lower for word in rate_role.split()[:3]):
                    daily_rate = rate
                    break
            
            if not daily_rate:
                # Tarifa por defecto según nivel
                if "senior" in profile_lower or "arquitecto" in profile_lower or "líder" in profile_lower:
                    daily_rate = 110000
                elif "junior" in profile_lower:
                    daily_rate = 55000
                else:
                    daily_rate = 85000
            
            role_cost = daily_rate * days
            total_cost += role_cost
            total_days += days
            
            line_items.append({
                "role": profile_name,
                "seniority": "Senior" if "senior" in profile_lower or "arquitecto" in profile_lower else "Mid",
                "quantity": 1,  # En Gantt cada asignación es de 1 persona
                "daily_rate": daily_rate,
                "days": days,
                "subtotal": role_cost
            })
    
    # Fallback: usar team_data si no hay datos de Gantt
    elif request.team_data and request.team_data.get("team_members"):
        # Usar duración total del Gantt o default
        project_days = 60
        if request.gantt_data:
            total_weeks = request.gantt_data.get("total_weeks", 12)
            project_days = total_weeks * 5
        
        for member in request.team_data["team_members"]:
            role = member.get("role", "").lower().strip()
            quantity = member.get("quantity", 1)
            seniority = member.get("seniority", "").lower()
            
            daily_rate = None
            if role in DAILY_RATES_COP:
                daily_rate = DAILY_RATES_COP[role]
            else:
                for rate_role, rate in DAILY_RATES_COP.items():
                    if role in rate_role or rate_role in role:
                        daily_rate = rate
                        break
            
            if not daily_rate:
                daily_rate = 85000
            
            role_cost = daily_rate * project_days * quantity
            total_cost += role_cost
            total_days = project_days
            
            line_items.append({
                "role": member.get("role", "Rol no especificado"),
                "seniority": member.get("seniority", "N/A"),
                "quantity": quantity,
                "daily_rate": daily_rate,
                "days": project_days,
                "subtotal": role_cost
            })
    
    # Si no hay ningún dato, generar estimación mínima
    if not line_items:
        default_team = [
            ("Project Manager", 1, 139341, 60),
            ("Desarrollador Senior", 2, 92746, 60),
            ("Desarrollador Junior", 2, 50489, 60),
            ("Tester QA", 1, 53419, 40),
        ]
        for role, qty, rate, days in default_team:
            cost = rate * days * qty
            total_cost += cost
            line_items.append({
                "role": role,
                "seniority": "Senior" if "senior" in role.lower() else "Mid",
                "quantity": qty,
                "daily_rate": rate,
                "days": days,
                "subtotal": cost
            })
        total_days = 60
    
    return EconomicsResponse(
        line_items=line_items,
        total_cost=total_cost,
        total_days=total_days,
        summary=f"Presupuesto basado en participación de {len(line_items)} perfiles: ${total_cost:,.0f} COP"
    )


# ============================================================
# RISKS ANALYSIS ENDPOINT
# ============================================================

class RisksRequest(BaseModel):
    documentation: str
    project_name: str = "Proyecto"
    architecture_data: Optional[dict] = None
    epics_data: Optional[dict] = None
    team_data: Optional[dict] = None
    gantt_data: Optional[dict] = None

class RisksResponse(BaseModel):
    artifact_type: str = "risks"
    risks: List[dict]
    total_risks: int
    high_priority_count: int
    markdown_content: str
    summary: str

@app.post("/api/artifact/risks", response_model=RisksResponse)
async def generate_risks(request: RisksRequest):
    """Genera análisis de riesgos del ciclo de vida de desarrollo usando LLM"""
    
    if not request.documentation or len(request.documentation.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="La documentación debe tener contenido suficiente"
        )
    
    try:
        from langchain_openai import AzureChatOpenAI
        
        llm = AzureChatOpenAI(
            azure_deployment="gpt-4.1",
            openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2025-01-01-preview"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            temperature=0.4,
            max_tokens=4000
        )
        
        # Construir contexto con datos de otros artefactos
        context_parts = [f"DOCUMENTACIÓN DEL PROYECTO:\n{request.documentation[:3000]}"]
        
        if request.architecture_data:
            arch_desc = request.architecture_data.get("description", "")[:1000]
            context_parts.append(f"\nARQUITECTURA PROPUESTA:\n{arch_desc}")
        
        if request.epics_data:
            epics_md = request.epics_data.get("markdown_content", "")[:1500]
            context_parts.append(f"\nÉPICAS DEL PROYECTO:\n{epics_md}")
        
        if request.team_data:
            team = request.team_data.get("team_members", [])
            team_summary = ", ".join([f"{m.get('role')} ({m.get('quantity')})" for m in team[:8]])
            context_parts.append(f"\nEQUIPO PROPUESTO: {team_summary}")
        
        if request.gantt_data:
            weeks = request.gantt_data.get("total_weeks", 0)
            tasks = len(request.gantt_data.get("tasks", []))
            context_parts.append(f"\nCRONOGRAMA: {weeks} semanas, {tasks} tareas")
        
        context = "\n".join(context_parts)
        
        prompt = f"""Eres un experto en gestión de riesgos de proyectos de software. Analiza el siguiente proyecto e identifica TODOS los riesgos potenciales en el ciclo de vida de desarrollo.

{context}

INSTRUCCIONES:
1. Identifica riesgos en TODAS las fases del SDLC:
   - Análisis y Requisitos
   - Diseño y Arquitectura
   - Desarrollo/Implementación
   - Pruebas/QA
   - Despliegue/DevOps
   - Mantenimiento/Operación

2. Para cada riesgo incluye:
   - Categoría (Técnico, Organizacional, Externo, Cronograma, Presupuesto)
   - Fase del SDLC afectada
   - Probabilidad (Alta, Media, Baja)
   - Impacto (Alto, Medio, Bajo)
   - Descripción del riesgo
   - Plan de mitigación

3. Prioriza los riesgos más críticos (Alta probabilidad + Alto impacto)

Responde SOLO con JSON válido:
{{
    "risks": [
        {{
            "id": "R001",
            "category": "Técnico",
            "phase": "Desarrollo",
            "name": "Nombre corto del riesgo",
            "description": "Descripción detallada del riesgo",
            "probability": "Alta",
            "impact": "Alto",
            "priority": "Crítica",
            "mitigation": "Plan de mitigación sugerido"
        }},
        ...
    ],
    "summary": "Resumen ejecutivo del análisis de riesgos"
}}"""
        
        response = llm.invoke(prompt)
        content = response.content.strip()
        
        # Extraer JSON
        import re
        json_match = re.search(r'\{[\s\S]*?"risks"[\s\S]*?\][\s\S]*?\}', content)
        
        if json_match:
            json_str = json_match.group()
            data = json.loads(json_str)
            
            risks = data.get("risks", [])
            summary = data.get("summary", "Análisis de riesgos completado")
            
            # Contar riesgos de alta prioridad
            high_priority = sum(1 for r in risks if r.get("priority", "").lower() in ["crítica", "alta", "critical", "high"])
            
            # Generar markdown
            markdown_lines = [f"# Análisis de Riesgos - {request.project_name}\n"]
            markdown_lines.append(f"**Total de riesgos identificados:** {len(risks)}\n")
            markdown_lines.append(f"**Riesgos de alta prioridad:** {high_priority}\n")
            markdown_lines.append("---\n")
            
            # Agrupar por categoría
            categories = {}
            for risk in risks:
                cat = risk.get("category", "Otro")
                if cat not in categories:
                    categories[cat] = []
                categories[cat].append(risk)
            
            for cat, cat_risks in categories.items():
                markdown_lines.append(f"\n## {cat}\n")
                for risk in cat_risks:
                    priority_emoji = "🔴" if risk.get("priority", "").lower() in ["crítica", "alta"] else "🟡" if risk.get("priority", "").lower() == "media" else "🟢"
                    markdown_lines.append(f"\n### {priority_emoji} {risk.get('id', 'R')} - {risk.get('name', 'Riesgo')}")
                    markdown_lines.append(f"- **Fase:** {risk.get('phase', 'N/A')}")
                    markdown_lines.append(f"- **Probabilidad:** {risk.get('probability', 'N/A')}")
                    markdown_lines.append(f"- **Impacto:** {risk.get('impact', 'N/A')}")
                    markdown_lines.append(f"- **Descripción:** {risk.get('description', '')}")
                    markdown_lines.append(f"- **Mitigación:** {risk.get('mitigation', '')}")
            
            markdown_content = "\n".join(markdown_lines)
            
            return RisksResponse(
                risks=risks,
                total_risks=len(risks),
                high_priority_count=high_priority,
                markdown_content=markdown_content,
                summary=summary
            )
        else:
            raise HTTPException(status_code=500, detail="Error al procesar respuesta del LLM")
            
    except Exception as e:
        logger.error(f"Error en análisis de riesgos: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


class DownloadRisksRequest(BaseModel):
    markdown_content: str
    project_name: str = "Proyecto"


@app.post("/api/artifact/risks/download")
async def download_risks_word(request: DownloadRisksRequest):
    """Descarga el análisis de riesgos como documento Word"""
    import io
    
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=500, detail="python-docx no instalado")
    
    try:
        doc = Document()
        
        # Título
        title = doc.add_heading('Análisis de Riesgos', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Procesar Markdown
        lines = request.markdown_content.split('\n')
        
        for line in lines:
            try:
                line = line.strip()
                if not line or line == '---':
                    continue
                
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- **') or line.startswith('- '):
                    add_markdown_paragraph(doc, line[2:], style='List Bullet')
                elif line.startswith('**'):
                    add_markdown_paragraph(doc, line)
                else:
                    add_markdown_paragraph(doc, line)
            except:
                continue
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        safe_name = request.project_name.replace(' ', '_').replace('/', '_')
        
        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=riesgos_{safe_name}.docx"
            }
        )
        
    except Exception as e:
        logger.error(f"Error generando Word de riesgos: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


# ============================================================
# PREREQUISITES ENDPOINT
# ============================================================

class PrerequisitesRequest(BaseModel):
    documentation: str
    project_name: str = "Proyecto"
    architecture_data: Optional[dict] = None
    epics_data: Optional[dict] = None
    team_data: Optional[dict] = None

class PrerequisitesResponse(BaseModel):
    artifact_type: str = "prerequisites"
    prerequisites: List[dict]
    total_count: int
    critical_count: int
    markdown_content: str
    summary: str

@app.post("/api/artifact/prerequisites", response_model=PrerequisitesResponse)
async def generate_prerequisites(request: PrerequisitesRequest):
    """Genera lista de prerrequisitos necesarios antes de iniciar desarrollo"""
    
    if not request.documentation or len(request.documentation.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="La documentación debe tener contenido suficiente"
        )
    
    try:
        from langchain_openai import AzureChatOpenAI
        
        llm = AzureChatOpenAI(
            azure_deployment="gpt-4.1",
            openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2025-01-01-preview"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            temperature=0.3,
            max_tokens=4000
        )
        
        # Construir contexto
        context_parts = [f"DOCUMENTACIÓN DEL PROYECTO:\n{request.documentation[:3000]}"]
        
        if request.architecture_data:
            arch_desc = request.architecture_data.get("description", "")[:1000]
            context_parts.append(f"\nARQUITECTURA PROPUESTA:\n{arch_desc}")
        
        if request.epics_data:
            epics_md = request.epics_data.get("markdown_content", "")[:1500]
            context_parts.append(f"\nÉPICAS DEL PROYECTO:\n{epics_md}")
        
        if request.team_data:
            team = request.team_data.get("team_members", [])
            team_summary = ", ".join([f"{m.get('role')} ({m.get('quantity')})" for m in team[:8]])
            context_parts.append(f"\nEQUIPO PROPUESTO: {team_summary}")
        
        context = "\n".join(context_parts)
        
        prompt = f"""Eres un experto en gestión de proyectos de software. Analiza el siguiente proyecto e identifica TODOS los prerrequisitos que deben estar disponibles ANTES de iniciar el desarrollo para evitar bloqueos.

{context}

INSTRUCCIONES:
1. Identifica prerrequisitos en las siguientes categorías:
   - **Infraestructura**: Servidores, ambientes, cloud, redes
   - **Accesos y Permisos**: Credenciales, VPN, repositorios, bases de datos
   - **Herramientas**: Licencias de software, IDEs, herramientas de CI/CD
   - **Documentación**: Specs técnicos, APIs existentes, diagramas
   - **Integraciones**: APIs de terceros, servicios externos, SSO
   - **Datos**: Bases de datos, datos de prueba, migraciones
   - **Recursos Humanos**: Capacitaciones, onboarding, asignaciones
   - **Aprobaciones**: Sign-offs, contratos, NDAs, compliance

2. Para cada prerrequisito incluye:
   - Categoría
   - Nombre corto
   - Descripción detallada
   - Criticidad (Crítico, Importante, Deseable)
   - Responsable sugerido (área o rol)
   - Tiempo estimado para obtenerlo

3. Prioriza los prerrequisitos CRÍTICOS que bloquearían el inicio

Responde SOLO con JSON válido:
{{
    "prerequisites": [
        {{
            "id": "P001",
            "category": "Infraestructura",
            "name": "Ambiente de desarrollo",
            "description": "Configuración del ambiente de desarrollo con todas las dependencias",
            "criticality": "Crítico",
            "responsible": "DevOps",
            "estimated_time": "3 días"
        }},
        ...
    ],
    "summary": "Resumen ejecutivo de los prerrequisitos identificados"
}}"""
        
        response = llm.invoke(prompt)
        content = response.content.strip()
        
        # Extraer JSON
        import re
        json_match = re.search(r'\{[\s\S]*?"prerequisites"[\s\S]*?\][\s\S]*?\}', content)
        
        if json_match:
            json_str = json_match.group()
            data = json.loads(json_str)
            
            prerequisites = data.get("prerequisites", [])
            summary = data.get("summary", "Prerrequisitos identificados")
            
            # Contar críticos
            critical_count = sum(1 for p in prerequisites if p.get("criticality", "").lower() in ["crítico", "critical"])
            
            # Generar markdown
            markdown_lines = [f"# Prerrequisitos - {request.project_name}\n"]
            markdown_lines.append(f"**Total de prerrequisitos:** {len(prerequisites)}\n")
            markdown_lines.append(f"**Prerrequisitos críticos:** {critical_count}\n")
            markdown_lines.append("---\n")
            
            # Agrupar por categoría
            categories = {}
            for prereq in prerequisites:
                cat = prereq.get("category", "Otro")
                if cat not in categories:
                    categories[cat] = []
                categories[cat].append(prereq)
            
            for cat, cat_prereqs in categories.items():
                markdown_lines.append(f"\n## {cat}\n")
                for prereq in cat_prereqs:
                    crit_emoji = "🔴" if prereq.get("criticality", "").lower() in ["crítico", "critical"] else "🟡" if prereq.get("criticality", "").lower() == "importante" else "🟢"
                    markdown_lines.append(f"\n### {crit_emoji} {prereq.get('id', 'P')} - {prereq.get('name', 'Prerrequisito')}")
                    markdown_lines.append(f"- **Criticidad:** {prereq.get('criticality', 'N/A')}")
                    markdown_lines.append(f"- **Responsable:** {prereq.get('responsible', 'N/A')}")
                    markdown_lines.append(f"- **Tiempo estimado:** {prereq.get('estimated_time', 'N/A')}")
                    markdown_lines.append(f"- **Descripción:** {prereq.get('description', '')}")
            
            markdown_content = "\n".join(markdown_lines)
            
            return PrerequisitesResponse(
                prerequisites=prerequisites,
                total_count=len(prerequisites),
                critical_count=critical_count,
                markdown_content=markdown_content,
                summary=summary
            )
        else:
            raise HTTPException(status_code=500, detail="Error al procesar respuesta del LLM")
            
    except Exception as e:
        logger.error(f"Error en prerrequisitos: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


class DownloadPrerequisitesRequest(BaseModel):
    markdown_content: str
    project_name: str = "Proyecto"


@app.post("/api/artifact/prerequisites/download")
async def download_prerequisites_word(request: DownloadPrerequisitesRequest):
    """Descarga los prerrequisitos como documento Word"""
    import io
    
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=500, detail="python-docx no instalado")
    
    try:
        doc = Document()
        
        title = doc.add_heading('Prerrequisitos del Proyecto', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        lines = request.markdown_content.split('\n')
        
        for line in lines:
            try:
                line = line.strip()
                if not line or line == '---':
                    continue
                
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- **'):
                    text = line[2:].replace('**', '')
                    doc.add_paragraph(text, style='List Bullet')
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('**'):
                    p = doc.add_paragraph()
                    text = line.replace('**', '')
                    run = p.add_run(text)
                    run.bold = True
                else:
                    doc.add_paragraph(line)
            except:
                continue
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        safe_name = request.project_name.replace(' ', '_').replace('/', '_')
        
        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=prerrequisitos_{safe_name}.docx"
            }
        )
        
    except Exception as e:
        logger.error(f"Error generando Word de prerrequisitos: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


# ============================================================
# CLIENT QUESTIONS ENDPOINT (Consultas al Cliente)
# ============================================================

class QuestionsRequest(BaseModel):
    documentation: str
    project_name: str = "Proyecto"
    architecture_data: Optional[dict] = None
    epics_data: Optional[dict] = None

class QuestionsResponse(BaseModel):
    artifact_type: str = "questions"
    questions: List[dict]
    total_count: int
    functional_count: int
    technical_count: int
    markdown_content: str
    summary: str

@app.post("/api/artifact/questions", response_model=QuestionsResponse)
async def generate_questions(request: QuestionsRequest):
    """Genera preguntas funcionales y técnicas para el cliente"""
    
    if not request.documentation or len(request.documentation.strip()) < 50:
        raise HTTPException(status_code=400, detail="Documentación insuficiente")
    
    try:
        from langchain_openai import AzureChatOpenAI
        
        llm = AzureChatOpenAI(
            azure_deployment="gpt-4.1",
            openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2025-01-01-preview"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            temperature=0.4,
            max_tokens=4000
        )
        
        context_parts = [f"DOCUMENTACIÓN:\n{request.documentation[:3500]}"]
        if request.architecture_data:
            context_parts.append(f"\nARQUITECTURA:\n{request.architecture_data.get('description', '')[:1000]}")
        if request.epics_data:
            context_parts.append(f"\nÉPICAS:\n{request.epics_data.get('markdown_content', '')[:1500]}")
        
        context = "\n".join(context_parts)
        
        prompt = f"""Eres un analista de negocios experto. Analiza la documentación y genera PREGUNTAS importantes que se deben hacer al cliente para aclarar dudas antes de iniciar el desarrollo.

{context}

INSTRUCCIONES:
1. Identifica vacíos de información en la documentación
2. Genera preguntas en estas categorías:
   - **Funcionales**: Reglas de negocio, flujos, validaciones, casos de uso
   - **Técnicas**: Integraciones, APIs, seguridad, rendimiento, infraestructura
   - **Datos**: Modelos, migraciones, volúmenes, formatos
   - **UX/UI**: Diseños, interacciones, responsive, accesibilidad
   - **Negocio**: Prioridades, deadlines, stakeholders, alcance

3. Para cada pregunta incluye:
   - Tipo (Funcional/Técnica)
   - Categoría específica
   - La pregunta clara y concisa
   - Contexto de por qué es importante

Responde SOLO con JSON válido:
{{
    "questions": [
        {{
            "id": "Q001",
            "type": "Funcional",
            "category": "Reglas de Negocio",
            "question": "¿Cuál es el flujo cuando un usuario...?",
            "context": "La documentación no especifica este caso"
        }},
        ...
    ],
    "summary": "Resumen de las áreas con mayor necesidad de clarificación"
}}"""
        
        response = llm.invoke(prompt)
        content = response.content.strip()
        
        import re
        json_match = re.search(r'\{[\s\S]*?"questions"[\s\S]*?\][\s\S]*?\}', content)
        
        if json_match:
            data = json.loads(json_match.group())
            questions = data.get("questions", [])
            summary = data.get("summary", "Consultas identificadas")
            
            functional = sum(1 for q in questions if q.get("type", "").lower() == "funcional")
            technical = len(questions) - functional
            
            # Generar markdown
            md_lines = [f"# Consultas al Cliente - {request.project_name}\n"]
            md_lines.append(f"**Total:** {len(questions)} | **Funcionales:** {functional} | **Técnicas:** {technical}\n---\n")
            
            categories = {}
            for q in questions:
                cat = q.get("category", "Otro")
                if cat not in categories:
                    categories[cat] = []
                categories[cat].append(q)
            
            for cat, cat_qs in categories.items():
                md_lines.append(f"\n## {cat}\n")
                for q in cat_qs:
                    emoji = "📋" if q.get("type") == "Funcional" else "⚙️"
                    md_lines.append(f"\n### {emoji} {q.get('id', 'Q')} - {q.get('question', '')}")
                    md_lines.append(f"- **Tipo:** {q.get('type', 'N/A')}")
                    md_lines.append(f"- **Contexto:** {q.get('context', '')}")
            
            return QuestionsResponse(
                questions=questions,
                total_count=len(questions),
                functional_count=functional,
                technical_count=technical,
                markdown_content="\n".join(md_lines),
                summary=summary
            )
        else:
            raise HTTPException(status_code=500, detail="Error procesando respuesta LLM")
            
    except Exception as e:
        logger.error(f"Error en preguntas: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


class DownloadQuestionsRequest(BaseModel):
    markdown_content: str
    project_name: str = "Proyecto"

@app.post("/api/artifact/questions/download")
async def download_questions_word(request: DownloadQuestionsRequest):
    """Descarga consultas como Word"""
    import io
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=500, detail="python-docx no instalado")
    
    try:
        doc = Document()
        doc.add_heading('Consultas al Cliente', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for line in request.markdown_content.split('\n'):
            line = line.strip()
            if not line or line == '---':
                continue
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- **'):
                doc.add_paragraph(line[2:].replace('**', ''), style='List Bullet')
            elif line.startswith('**'):
                p = doc.add_paragraph()
                p.add_run(line.replace('**', '')).bold = True
            else:
                doc.add_paragraph(line)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=consultas_{request.project_name.replace(' ', '_')}.docx"}
        )
    except Exception as e:
        logger.error(f"Error Word consultas: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


# ============================================================
# TECHNICAL IMPROVEMENTS ENDPOINT (Mejoras Técnicas)
# ============================================================

class ImprovementsRequest(BaseModel):
    documentation: str
    project_name: str = "Proyecto"
    architecture_data: Optional[dict] = None
    epics_data: Optional[dict] = None
    team_data: Optional[dict] = None

class ImprovementsResponse(BaseModel):
    artifact_type: str = "improvements"
    improvements: List[dict]
    total_count: int
    high_impact_count: int
    markdown_content: str
    summary: str

@app.post("/api/artifact/improvements", response_model=ImprovementsResponse)
async def generate_improvements(request: ImprovementsRequest):
    """Genera sugerencias de mejoras técnicas para el proyecto"""
    
    if not request.documentation or len(request.documentation.strip()) < 50:
        raise HTTPException(status_code=400, detail="Documentación insuficiente")
    
    try:
        from langchain_openai import AzureChatOpenAI
        
        llm = AzureChatOpenAI(
            azure_deployment="gpt-4.1",
            openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2025-01-01-preview"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            temperature=0.4,
            max_tokens=4000
        )
        
        context_parts = [f"DOCUMENTACIÓN:\n{request.documentation[:3500]}"]
        if request.architecture_data:
            context_parts.append(f"\nARQUITECTURA:\n{request.architecture_data.get('description', '')[:1000]}")
        if request.epics_data:
            context_parts.append(f"\nÉPICAS:\n{request.epics_data.get('markdown_content', '')[:1500]}")
        
        context = "\n".join(context_parts)
        
        prompt = f"""Eres un arquitecto de software senior. Analiza el proyecto y sugiere MEJORAS TÉCNICAS que harían el desarrollo más robusto, escalable y mantenible.

{context}

INSTRUCCIONES:
1. Identifica oportunidades de mejora en:
   - **Arquitectura**: Patrones, microservicios, modularidad
   - **Seguridad**: Autenticación, autorización, encriptación, OWASP
   - **Rendimiento**: Caché, optimización, lazy loading, CDN
   - **Escalabilidad**: Load balancing, horizontal scaling, cloud-native
   - **DevOps**: CI/CD, containerización, IaC, monitoreo
   - **Calidad**: Testing, code review, linting, documentación
   - **UX**: Accesibilidad, PWA, offline-first, responsive

2. Para cada mejora incluye:
   - Categoría
   - Nombre de la mejora
   - Descripción detallada
   - Impacto (Alto, Medio, Bajo)
   - Esfuerzo estimado (Alto, Medio, Bajo)
   - Beneficio principal

Responde SOLO con JSON válido:
{{
    "improvements": [
        {{
            "id": "M001",
            "category": "Seguridad",
            "name": "Implementar OAuth 2.0",
            "description": "Usar OAuth 2.0 con JWT para autenticación",
            "impact": "Alto",
            "effort": "Medio",
            "benefit": "Mayor seguridad y estándar de la industria"
        }},
        ...
    ],
    "summary": "Resumen de las mejoras más importantes"
}}"""
        
        response = llm.invoke(prompt)
        content = response.content.strip()
        
        import re
        json_match = re.search(r'\{[\s\S]*?"improvements"[\s\S]*?\][\s\S]*?\}', content)
        
        if json_match:
            data = json.loads(json_match.group())
            improvements = data.get("improvements", [])
            summary = data.get("summary", "Mejoras identificadas")
            
            high_impact = sum(1 for m in improvements if m.get("impact", "").lower() == "alto")
            
            # Generar markdown
            md_lines = [f"# Mejoras Técnicas - {request.project_name}\n"]
            md_lines.append(f"**Total:** {len(improvements)} | **Alto Impacto:** {high_impact}\n---\n")
            
            categories = {}
            for m in improvements:
                cat = m.get("category", "Otro")
                if cat not in categories:
                    categories[cat] = []
                categories[cat].append(m)
            
            for cat, cat_ms in categories.items():
                md_lines.append(f"\n## {cat}\n")
                for m in cat_ms:
                    impact_emoji = "🔴" if m.get("impact") == "Alto" else "🟡" if m.get("impact") == "Medio" else "🟢"
                    md_lines.append(f"\n### {impact_emoji} {m.get('id', 'M')} - {m.get('name', 'Mejora')}")
                    md_lines.append(f"- **Impacto:** {m.get('impact', 'N/A')} | **Esfuerzo:** {m.get('effort', 'N/A')}")
                    md_lines.append(f"- **Descripción:** {m.get('description', '')}")
                    md_lines.append(f"- **Beneficio:** {m.get('benefit', '')}")
            
            return ImprovementsResponse(
                improvements=improvements,
                total_count=len(improvements),
                high_impact_count=high_impact,
                markdown_content="\n".join(md_lines),
                summary=summary
            )
        else:
            raise HTTPException(status_code=500, detail="Error procesando respuesta LLM")
            
    except Exception as e:
        logger.error(f"Error en mejoras: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


class DownloadImprovementsRequest(BaseModel):
    markdown_content: str
    project_name: str = "Proyecto"

@app.post("/api/artifact/improvements/download")
async def download_improvements_word(request: DownloadImprovementsRequest):
    """Descarga mejoras técnicas como Word"""
    import io
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=500, detail="python-docx no instalado")
    
    try:
        doc = Document()
        doc.add_heading('Mejoras Técnicas', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for line in request.markdown_content.split('\n'):
            line = line.strip()
            if not line or line == '---':
                continue
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- **'):
                doc.add_paragraph(line[2:].replace('**', ''), style='List Bullet')
            elif line.startswith('**'):
                p = doc.add_paragraph()
                p.add_run(line.replace('**', '')).bold = True
            else:
                doc.add_paragraph(line)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=mejoras_{request.project_name.replace(' ', '_')}.docx"}
        )
    except Exception as e:
        logger.error(f"Error Word mejoras: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")

# ============================================================
# CONSOLIDATED EXPORT ENDPOINT
# ============================================================

class ExportAllRequest(BaseModel):
    project_name: str = "Proyecto"
    artifacts: dict

@app.post("/api/artifact/export-all")
async def export_all_word(request: ExportAllRequest):
    """Genera un documento Word consolidado con todos los artefactos"""
    import io
    
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=500, detail="python-docx no instalado")
    
    try:
        doc = Document()
        
        # Portada
        title = doc.add_heading(f'Propuesta Técnica', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(request.project_name)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Índice de contenido
        doc.add_heading('Contenido', level=1)
        artifacts_order = ['architecture', 'team', 'epics', 'gantt', 'economics', 'risks', 'prerequisites', 'questions', 'improvements']
        artifact_titles = {
            'architecture': '1. Arquitectura del Sistema',
            'team': '2. Propuesta de Equipo',
            'epics': '3. Épicas del Proyecto',
            'gantt': '4. Cronograma (Gantt)',
            'economics': '5. Propuesta Económica',
            'risks': '6. Análisis de Riesgos',
            'prerequisites': '7. Prerrequisitos',
            'questions': '8. Consultas al Cliente',
            'improvements': '9. Mejoras Técnicas'
        }
        
        for key in artifacts_order:
            if key in request.artifacts:
                doc.add_paragraph(artifact_titles.get(key, key), style='List Number')
        
        doc.add_page_break()
        
        # Contenido de cada artefacto
        for key in artifacts_order:
            if key not in request.artifacts:
                continue
                
            artifact = request.artifacts[key]
            doc.add_heading(artifact_titles.get(key, key), level=1)
            
            # Descripción/Resumen
            summary = artifact.get('summary', artifact.get('description', ''))
            if summary:
                doc.add_paragraph(summary)
            
            # Markdown content (epics, questions, improvements, etc.)
            md_content = artifact.get('markdown_content', '')
            if md_content:
                for line in md_content.split('\n'):
                    original_line = line
                    line = line.strip()
                    if not line or line == '---':
                        continue
                    # Headings
                    if line.startswith('#### '):
                        doc.add_heading(line[5:], level=4)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('# '):
                        doc.add_heading(line[2:], level=2)
                    # Checkboxes
                    elif line.startswith('- [ ] ') or line.startswith('- [x] '):
                        checked = '☑' if '[x]' in line else '☐'
                        text = line[6:]
                        doc.add_paragraph(f"{checked} {text}", style='List Bullet')
                    # Numbered lists
                    elif len(line) > 2 and line[0].isdigit() and line[1] in '.):':
                        doc.add_paragraph(line[2:].strip(), style='List Number')
                    elif len(line) > 3 and line[:2].isdigit() and line[2] in '.):':
                        doc.add_paragraph(line[3:].strip(), style='List Number')
                    # Bullet lists
                    elif line.startswith('- ') or line.startswith('* '):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    else:
                        p = doc.add_paragraph()
                        # Handle bold and italic text
                        remaining = line
                        while remaining:
                            if '**' in remaining:
                                before, _, after = remaining.partition('**')
                                if before:
                                    p.add_run(before)
                                if '**' in after:
                                    bold_text, _, remaining = after.partition('**')
                                    run = p.add_run(bold_text)
                                    run.bold = True
                                else:
                                    remaining = after
                            else:
                                p.add_run(remaining)
                                break
            
            # Team members
            team = artifact.get('team_members', [])
            if team:
                for member in team:
                    role = member.get('role', 'Rol')
                    qty = member.get('quantity', 1)
                    doc.add_paragraph(f"• {role} x{qty}", style='List Bullet')
            
            # Economics
            if key == 'economics':
                items = artifact.get('line_items', [])
                total = artifact.get('total_cost', 0)
                for item in items:
                    profile = item.get('profile', '')
                    cost = item.get('cost', 0)
                    doc.add_paragraph(f"• {profile}: ${cost:,.0f} COP", style='List Bullet')
                doc.add_paragraph(f"Total: ${total:,.0f} COP")
            
            # Gantt tasks
            if key == 'gantt':
                tasks = artifact.get('tasks', [])
                for task in tasks[:15]:  # Limit to avoid huge docs
                    name = task.get('name', '')
                    start = task.get('start_date', '')
                    end = task.get('end_date', '')
                    doc.add_paragraph(f"• {name} ({start} - {end})", style='List Bullet')
            
            doc.add_page_break()
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        safe_name = request.project_name.replace(' ', '_').replace('/', '_')
        
        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=propuesta_completa_{safe_name}.docx"
            }
        )
        
    except Exception as e:
        logger.error(f"Error generando Word consolidado: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "orchestrator_loaded": HAS_ORCHESTRATOR,
        "timestamp": datetime.now().isoformat()
    }


@app.get("/")
async def root():
    """Redirecciona a index.html"""
    return FileResponse(Path(__file__).parent / "web" / "index.html")


# ============================================================
# SERVIR ARCHIVOS ESTÁTICOS
# ============================================================

web_dir = Path(__file__).parent / "web"
if web_dir.exists():
    app.mount("/web", StaticFiles(directory=web_dir), name="static")


if __name__ == "__main__":
    import uvicorn
    print("\n🚀 Iniciando Project Generator Agent...")
    print("📍 URL: http://localhost:8011")
    print("📖 Docs: http://localhost:8011/docs\n")
    uvicorn.run(app, host="0.0.0.0", port=8011)
